import logging
from pathlib import Path
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.controllers.base_controller import BaseController
from app.core.config import Settings
from app.models.contract import Contract
from PyPDF2 import PdfReader
from docx import Document
import re
from app.models.chunk import Chunk
from app.enums import ProcessingEnums, ChunkEnums
from app.controllers.cohere_controller import CohereController
from app.controllers.qdrant_controller import QdrantController
from app.controllers.groq_controller import GroqController
from app.models.risk_score import RiskScore
import nltk
from nltk.tokenize import sent_tokenize

logger = logging.getLogger(__name__)

class ProcessingController(BaseController):
    def __init__(self, db: AsyncSession, settings: Settings):
        super().__init__(db, settings)

        self.cohere = CohereController(db=db, settings=settings)
        self.qdrant = QdrantController(db=db, settings=settings)
        self.groq = GroqController(db=db, settings=settings)

    async def process(self, contract_id: str):
        result = await self.db.execute(select(Contract).where(Contract.uuid == contract_id))
        contract = result.scalar_one_or_none()
        if not contract:
            return False, "Contract not found"

        await self._update_status(contract, ProcessingEnums.PROCESSING)

        text = self._process_text(contract.file_path)
        chunks = await self._chunk_text(text, contract.id)

        await self._embed_and_store(chunks)
        await self._analyze(contract, text)
        await self._update_status(contract, ProcessingEnums.COMPLETED)

        return True, "Contract processed successfully"

    async def _update_status(self, contract: Contract, status: ProcessingEnums):
        contract.status = status
        await self.db.commit()
        await self.db.refresh(contract)

    def _process_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        logger.info(f"Extracting text from {ext} file")

        if ext == ".pdf":
            return self._process_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._process_docx(file_path)
        elif ext == ".txt":
            return self._process_txt(file_path)
        else:
            logger.error(f"Unsupported file type: {ext}")
            return ""

    def _process_pdf(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        logger.info(f"Extracted {len(reader.pages)} pages from PDF")
        return text

    def _process_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        logger.info(f"Extracted {len(doc.paragraphs)} paragraphs from DOCX")
        return text

    def _process_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        logger.info(f"Extracted text from TXT file")
        return text

    async def _chunk_text(self, text: str, contract_id: str) -> list:
        nltk.download('punkt', quiet=True)
        nltk.download('punkt_tab', quiet=True)

        chunks = []
        child_size = 100
        overlap_sentences = 2

        section_pattern = r'\n(?=(?:Section\s+\d+|SECTION\s+\d+|\d+\.\s+[A-Z]))'
        sections = re.split(section_pattern, text)
        sections = [s.strip() for s in sections if s.strip()]

        if len(sections) <= 1:
            sections = [text]

        for section in sections:
            lines = section.split('\n')
            section_title = lines[0].strip()
            section_text = '\n'.join(lines[1:]).strip() if len(lines) > 1 else section

            if not section_text:
                section_text = section

            parent = Chunk(
                contract_id=contract_id,
                chunk_type=ChunkEnums.PARENT,
                text=section_text,
                section_title=section_title,
            )
            self.db.add(parent)
            await self.db.flush()

            sentences = sent_tokenize(section_text)

            current_chunk = []
            current_word_count = 0
            child_count = 0

            for sentence in sentences:
                sentence_words = sentence.split()
                current_chunk.append(sentence)
                current_word_count += len(sentence_words)

                if current_word_count >= child_size:
                    child_text = " ".join(current_chunk)
                    child = Chunk(
                        contract_id=contract_id,
                        parent_id=parent.id,
                        chunk_type=ChunkEnums.CHILD,
                        text=child_text,
                        section_title=section_title,
                    )
                    self.db.add(child)
                    chunks.append(child)
                    child_count += 1

                    # overlap
                    current_chunk = current_chunk[-overlap_sentences:]
                    current_word_count = sum(len(s.split()) for s in current_chunk)

            # save remaining sentences as last chunk
            if current_chunk:
                remaining_text = " ".join(current_chunk)
                if len(remaining_text.split()) >= 10:
                    child = Chunk(
                        contract_id=contract_id,
                        parent_id=parent.id,
                        chunk_type=ChunkEnums.CHILD,
                        text=remaining_text,
                        section_title=section_title,
                    )
                    self.db.add(child)
                    chunks.append(child)
                    child_count += 1

        await self.db.commit()
        return chunks

    async def _embed_and_store(self, chunks: list):
        texts = [chunk.text for chunk in chunks]
        dense_vectors = self.cohere.embed_documents(texts)
        sparse_vectors = self.qdrant.embed_sparse(texts)
        self.qdrant.ensure_collection()
        self.qdrant.store_chunks(chunks, dense_vectors, sparse_vectors)

    async def _analyze(self, contract: Contract, text: str):
        result = await self.groq.analyze_contract(text)

        contract.summary = result.get("summary")
        contract.overall_risk_score = result.get("overall_risk_score")
        contract.contract_type = result.get("contract_type", "other")
        await self.db.commit()

        risk = RiskScore(
            contract_id=contract.id,
            overall_score=result.get("overall_risk_score", 0),
            ip_clauses_score=result.get("ip_clauses_score"),
            termination_score=result.get("termination_score"),
            non_compete_score=result.get("non_compete_score"),
            payment_score=result.get("payment_score"),
            auto_renewal_score=result.get("auto_renewal_score"),
            red_flags=str(result.get("red_flags", []))
        )
        self.db.add(risk)
        await self.db.commit()
        logger.info(f"Analysis completed for contract: {contract.id}")